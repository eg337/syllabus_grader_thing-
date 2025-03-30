from pypdf import PdfReader
from google import genai
from docx import Document
from bs4 import BeautifulSoup
import json
from lxml import etree

def parse_schedule(grades, file, name):
    with open("key.txt") as f:
        for line in f:
            API_KEY = line.strip()
   

    # name = 'CMSC470-Syllabus.pdf'
    # name = "syllabus_416_25.pdf"
    # name = "3220 SP25 Syllabus_UpdatedJan23.docx"
    # name = "Advanced Data Structures and Algorithms.html"
    # name = "STAT400-Syllabus.pdf"
    # name = "Adv. Algo. - Spring 2025 - Main - Google Drive.html"
    texts = ""
    tables = ""
    data = ""
    extension = name.split(".")[-1]

    if extension == "pdf":
        reader = PdfReader(file)
        for page in reader.pages:
            texts += page.extract_text()


    elif extension == "docx":
        document = Document(file)
        for paragraph in document.paragraphs:
            texts += paragraph.text
        
        data = []
        for index in range(len(document.tables)):
            table = document.tables[index]    
            keys = None
            for i, row in enumerate(table.rows):
                text = (cell.text for cell in row.cells)

                # Establish the mapping based on the first row
                # headers; these will become the keys of our dictionary
                if i == 0:
                    keys = tuple(text)
                    continue

                # Construct a dictionary for this row, mapping
                # keys to values for this row
                row_data = dict(zip(keys, text))
                data.append(row_data)
            # print("CHECK1")
            # print(data)
            # print("CHECK2")
    elif extension == "html":
        print(file)
        html_content = file.read()
        
        soup = BeautifulSoup(html_content, "html.parser")
        texts = soup.get_text()
 
        table = etree.HTML(html_content).find("body/table")
        if table is not None:
            rows = iter(table)
            headers = [col.text for col in next(rows)]
            for row in rows:
                values = [col.text for col in row]
                print(dict(zip(headers, values)))
        # print("CHECK1")
        # print(texts)
        # print("CHECk2")

    #print(texts)
    client = genai.Client(api_key=API_KEY)

    response = client.models.generate_content(
        model="gemini-2.0-flash", contents="In the message text after, I will upload a syllabus or schedule of a college course; the weights of the types of assignments in each course are given in the following dictionary" +str(grades)+ "; parse the"
        +" syllabus for a schedule of the course containing assignments and their deadlines;"
       + " for presentations list the earliest date which the presentation occurs as the deadline; for midterms and tests list the date which they occur as the deadline; format the deadlines as mm/dd; return a Json sorted by the assigment type "
       + "use the assignment types as keys and have the values be a list of subdictionaries containing the deadlines" + texts + str(data)
    )

    #print(tables)
    #print(response.text)

    string_data = response.text.strip("`")[4:]

    print(string_data)
    while string_data[-1] != "}":
        string_data = string_data[:-1]
    try:
        json_object = json.loads(string_data)
        return json_object
    except json.JSONDecodeError as e:
        print(f"Error decoding JSON: {e}")

