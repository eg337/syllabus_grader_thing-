from pypdf import PdfReader
from google import genai
from docx import Document
from bs4 import BeautifulSoup
import json
from schedule_parse import parse_schedule

with open("key.txt") as f:
    for line in f:
        API_KEY = line.strip()

# name = 'CMSC470-Syllabus.pdf'
# name = "syllabus_416_25.pdf"
name = "3220 SP25 Syllabus_UpdatedJan23.docx"
# name = "Advanced Data Structures and Algorithms.html"
texts = ""

extension = name.split(".")[1]

if extension == "pdf":
    reader = PdfReader(name)
    for page in reader.pages:
        texts += page.extract_text()


elif extension == "docx":
    document = Document(name)
    for paragraph in document.paragraphs:
        texts += paragraph.text

elif extension == "html":
    with open(name, "r") as file:
        html_content = file.read()
    
    soup = BeautifulSoup(html_content, "html.parser")
    texts = soup.get_text()


client = genai.Client(api_key=API_KEY)

response = client.models.generate_content(
    model="gemini-2.0-flash", contents="In the message text after, I will upload a syllabus to a college course. After I do so, please output the grade weights as percentages in a subdictionary labeled 'grade weights', separated by a space,"
    + " followed by a percent sign contained inside in a json file, format with each assignment as a key with the value being a subdictionary, inside list the weight, the number of assignments of the given type,"
    + "and the number of drops defaulting to 0 drops if none are listed; Additionally in a separate dictionary, with key 'course title', separated by a space, list the title of the course" + texts 
)


#print(response.text)

string_data = response.text.strip("`")[4:]


try:
    json_object = json.loads(string_data)
    print(json_object)
    print(type(json_object)) 
except json.JSONDecodeError as e:
    print(f"Error decoding JSON: {e}")

def parse_grades():
    with open("key.txt") as f:
        for line in f:
            API_KEY = line.strip()

    # name = 'CMSC470-Syllabus.pdf'
    # name = "syllabus_416_25.pdf"
    # name = "3220 SP25 Syllabus_UpdatedJan23.docx"
    name = "Advanced Data Structures and Algorithms.html"
    # name = "STAT400-Syllabus.pdf"x
    texts = ""

    extension = name.split(".")[1]

    if extension == "pdf":
        reader = PdfReader(name)
        for page in reader.pages:
            texts += page.extract_text()


    elif extension == "docx":
        document = Document(name)
        for paragraph in document.paragraphs:
            texts += paragraph.text

    elif extension == "html":
        with open(name, "r") as file:
            html_content = file.read()
        
        soup = BeautifulSoup(html_content, "html.parser")
        texts = soup.get_text()


    client = genai.Client(api_key=API_KEY)

    response = client.models.generate_content(
        model="gemini-2.0-flash", contents="In the message text after, I will upload a syllabus to a college course. After I do so, please output the grade weights as percentages in a subdictionary labeled 'grade weights', separated by a space,"
        + " followed by a percent sign contained inside in a json file, format with each assignment as a key with the value being a subdictionary, inside list the weight, the number of assignments of the given type,"
        + "and the number of drops defaulting to 0 drops if none are listed; Additionally in a separate dictionary, with key 'course title', separated by a space, list the title of the course" + texts 
    )


    #print(response.text)

    string_data = response.text.strip("`")[4:]


    try:
        json_object = json.loads(string_data)
        return json_object['grade weights']
    except json.JSONDecodeError as e:
        print(f"Error decoding JSON: {e}")
        return None

print(parse_schedule(parse_grades()))